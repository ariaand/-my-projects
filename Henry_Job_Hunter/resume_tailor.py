"""
resume_tailor.py — AI-powered resume tailoring + DOCX/PDF generation.

For each job above the score threshold:
  1. Calls Claude to rewrite the professional summary, reorder competencies,
     and sharpen bullets using ATS keywords from the job description.
  2. Generates a clean ATS-friendly DOCX file.
  3. Generates a matching PDF file.
  4. Generates a short application strategy note.
"""

import json
import logging
import os
import re
from typing import Dict, List

import anthropic
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors

logger = logging.getLogger(__name__)


# ── Prompt for Claude ─────────────────────────────────────────────────────────

TAILOR_PROMPT = """You are an expert resume writer and ATS optimization specialist.

I will give you:
1. A candidate profile (JSON) — this is their REAL experience. Do NOT fabricate anything.
2. A job description
3. The job title and company name

Your task:
- Rewrite the professional summary (3-4 sentences) specifically for this job using keywords from the job description.
- Reorder the core_competencies list so the most relevant skills appear first (keep all of them, just reorder).
- Rewrite the bullet points in each experience entry to naturally include ATS keywords from the job description — but ONLY when the rewrite is truthful and grounded in the original bullet.
- Do NOT add skills, tools, certifications, or experience that are not in the original profile.
- Keep bullets concise (one line each, action verb first).
- Return ONLY valid JSON in exactly this structure:

{
  "summary": "...",
  "core_competencies": ["...", "..."],
  "experience": [
    {
      "title": "...",
      "company": "...",
      "location": "...",
      "dates": "...",
      "bullets": ["...", "..."]
    }
  ],
  "keywords_used": ["keyword1", "keyword2"],
  "changes_made": "Brief explanation of what was changed and why."
}

CANDIDATE PROFILE:
{profile}

JOB TITLE: {job_title}
COMPANY: {company}

JOB DESCRIPTION:
{job_description}
"""


def tailor_resume_with_ai(
    profile: Dict,
    job: Dict,
    model: str,
    api_key: str
) -> Dict:
    """Call Claude to produce a tailored resume JSON."""
    client = anthropic.Anthropic(api_key=api_key)

    prompt = TAILOR_PROMPT.format(
        profile=json.dumps(profile, indent=2),
        job_title=job.get("title", ""),
        company=job.get("company", ""),
        job_description=job.get("description", "")[:4000],  # Trim very long JDs
    )

    message = client.messages.create(
        model=model,
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = message.content[0].text.strip()

    # Extract JSON even if Claude wraps it in markdown code fences
    json_match = re.search(r"\{[\s\S]*\}", raw)
    if not json_match:
        raise ValueError("Claude did not return valid JSON for resume tailoring.")

    return json.loads(json_match.group())


# ── DOCX generation ───────────────────────────────────────────────────────────

def _add_horizontal_rule(doc: Document) -> None:
    """Add a simple horizontal line to the DOCX."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_font(run, name="Calibri", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def generate_docx(profile: Dict, tailored: Dict, output_path: str) -> None:
    """
    Build a clean, ATS-friendly DOCX resume.
    No tables, no text boxes, no headers/footers — pure paragraphs.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ── Header ────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(profile["name"].upper())
    _set_font(name_run, size=18, bold=True, color=(31, 73, 125))

    contact_parts = [
        profile.get("phone", ""),
        profile.get("email", ""),
        profile.get("location", ""),
    ]
    if profile.get("linkedin"):
        contact_parts.append(profile["linkedin"])
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run("  |  ".join(p for p in contact_parts if p))
    _set_font(contact_run, size=9.5)

    _add_horizontal_rule(doc)

    # ── Professional Summary ──────────────────────────────────────────────────
    sec = doc.add_paragraph()
    sec_run = sec.add_run("PROFESSIONAL SUMMARY")
    _set_font(sec_run, size=10.5, bold=True, color=(31, 73, 125))

    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_after = Pt(4)
    sum_run = summary_para.add_run(tailored.get("summary", profile.get("summary_base", "")))
    _set_font(sum_run, size=10)

    _add_horizontal_rule(doc)

    # ── Core Competencies ─────────────────────────────────────────────────────
    sec2 = doc.add_paragraph()
    sec2_run = sec2.add_run("CORE COMPETENCIES")
    _set_font(sec2_run, size=10.5, bold=True, color=(31, 73, 125))

    competencies = tailored.get("core_competencies", profile.get("core_competencies", []))
    # 3-column layout using tab stops (simple, ATS-friendly)
    for i in range(0, len(competencies), 3):
        row = competencies[i:i+3]
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        cols = "   •  ".join(row)
        run = p.add_run(cols)
        _set_font(run, size=10)

    _add_horizontal_rule(doc)

    # ── Work Experience ───────────────────────────────────────────────────────
    sec3 = doc.add_paragraph()
    sec3_run = sec3.add_run("WORK EXPERIENCE")
    _set_font(sec3_run, size=10.5, bold=True, color=(31, 73, 125))

    experience = tailored.get("experience", profile.get("experience", []))
    for exp in experience:
        # Title + dates on same line
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(4)
        title_para.paragraph_format.space_after = Pt(1)
        t_run = title_para.add_run(f"{exp['title'].upper()}  —  {exp['company']}")
        _set_font(t_run, size=10.5, bold=True)
        title_para.add_run(f"  |  {exp.get('location', '')}  |  {exp.get('dates', '')}")

        for bullet in exp.get("bullets", []):
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(1)
            bp.paragraph_format.left_indent = Inches(0.2)
            b_run = bp.add_run(bullet)
            _set_font(b_run, size=10)

    _add_horizontal_rule(doc)

    # ── Certifications ────────────────────────────────────────────────────────
    sec4 = doc.add_paragraph()
    _set_font(sec4.add_run("CERTIFICATIONS"), size=10.5, bold=True, color=(31, 73, 125))

    for cert in profile.get("certifications", []):
        cp = doc.add_paragraph(style="List Bullet")
        cp.paragraph_format.space_after = Pt(1)
        _set_font(cp.add_run(cert), size=10)

    _add_horizontal_rule(doc)

    # ── Education ─────────────────────────────────────────────────────────────
    sec5 = doc.add_paragraph()
    _set_font(sec5.add_run("EDUCATION"), size=10.5, bold=True, color=(31, 73, 125))

    for edu in profile.get("education", []):
        ep = doc.add_paragraph()
        ep.paragraph_format.space_after = Pt(1)
        _set_font(ep.add_run(f"{edu['institution']} — {edu['degree']}"), size=10)

    # ── Languages ─────────────────────────────────────────────────────────────
    lang_para = doc.add_paragraph()
    lang_para.paragraph_format.space_before = Pt(4)
    _set_font(lang_para.add_run("LANGUAGES: "), size=10, bold=True)
    _set_font(lang_para.add_run(", ".join(profile.get("languages", []))), size=10)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    logger.info(f"DOCX saved: {output_path}")


# ── PDF generation ────────────────────────────────────────────────────────────

def generate_pdf(profile: Dict, tailored: Dict, output_path: str) -> None:
    """Generate a clean ATS-friendly PDF using reportlab."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    story = []
    styles = getSampleStyleSheet()

    BLUE = colors.HexColor("#1F497D")
    DARK = colors.HexColor("#222222")
    GRAY = colors.HexColor("#666666")

    name_style = ParagraphStyle("Name", fontSize=18, textColor=BLUE,
                                 alignment=TA_CENTER, spaceAfter=4, fontName="Helvetica-Bold")
    contact_style = ParagraphStyle("Contact", fontSize=9.5, textColor=DARK,
                                    alignment=TA_CENTER, spaceAfter=6)
    section_style = ParagraphStyle("Section", fontSize=10.5, textColor=BLUE,
                                    fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle("Body", fontSize=10, textColor=DARK,
                                 leading=14, spaceAfter=3)
    bullet_style = ParagraphStyle("Bullet", fontSize=10, textColor=DARK,
                                   leading=13, spaceAfter=2, leftIndent=14,
                                   bulletIndent=4)
    exp_title_style = ParagraphStyle("ExpTitle", fontSize=10.5, textColor=DARK,
                                      fontName="Helvetica-Bold", spaceBefore=6, spaceAfter=2)
    meta_style = ParagraphStyle("Meta", fontSize=9.5, textColor=GRAY, spaceAfter=3)

    def hr():
        return HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#999999"),
                          spaceAfter=4, spaceBefore=2)

    # Name
    story.append(Paragraph(profile["name"].upper(), name_style))

    # Contact
    contact_parts = [profile.get("phone", ""), profile.get("email", ""),
                     profile.get("location", "")]
    if profile.get("linkedin"):
        contact_parts.append(profile["linkedin"])
    story.append(Paragraph("  |  ".join(p for p in contact_parts if p), contact_style))
    story.append(hr())

    # Summary
    story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
    summary = tailored.get("summary", profile.get("summary_base", ""))
    story.append(Paragraph(summary, body_style))
    story.append(hr())

    # Competencies
    story.append(Paragraph("CORE COMPETENCIES", section_style))
    competencies = tailored.get("core_competencies", profile.get("core_competencies", []))
    for i in range(0, len(competencies), 3):
        row = "   •   ".join(competencies[i:i+3])
        story.append(Paragraph(f"• {row}", bullet_style))
    story.append(hr())

    # Experience
    story.append(Paragraph("WORK EXPERIENCE", section_style))
    for exp in tailored.get("experience", profile.get("experience", [])):
        story.append(Paragraph(
            f"{exp['title'].upper()} — {exp['company']}",
            exp_title_style
        ))
        story.append(Paragraph(
            f"{exp.get('location', '')}  |  {exp.get('dates', '')}",
            meta_style
        ))
        for bullet in exp.get("bullets", []):
            story.append(Paragraph(f"• {bullet}", bullet_style))
    story.append(hr())

    # Certifications
    story.append(Paragraph("CERTIFICATIONS", section_style))
    for cert in profile.get("certifications", []):
        story.append(Paragraph(f"• {cert}", bullet_style))
    story.append(hr())

    # Education
    story.append(Paragraph("EDUCATION", section_style))
    for edu in profile.get("education", []):
        story.append(Paragraph(f"{edu['institution']} — {edu['degree']}", body_style))

    # Languages
    story.append(Spacer(1, 4))
    langs = ", ".join(profile.get("languages", []))
    story.append(Paragraph(f"<b>LANGUAGES:</b> {langs}", body_style))

    doc.build(story)
    logger.info(f"PDF saved: {output_path}")


# ── Application strategy note ─────────────────────────────────────────────────

def generate_application_note(
    job: Dict,
    tailored: Dict,
    output_path: str,
) -> None:
    """Write a short text file explaining the application strategy for this job."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    lines = [
        f"APPLICATION STRATEGY NOTE",
        f"{'='*50}",
        f"Job Title  : {job.get('title', '')}",
        f"Company    : {job.get('company', '')}",
        f"Source     : {job.get('source', '')}",
        f"Match Score: {job.get('score', 0)}/100",
        f"Score Detail: {job.get('score_breakdown', '')}",
        f"Link       : {job.get('link', '')}",
        f"Pay Range  : {job.get('salary', 'Not listed')}",
        f"",
        f"WHY THIS IS A STRONG FIT",
        f"-" * 30,
        f"Software matched : {', '.join(job.get('software_found', [])) or 'None detected'}",
        f"Red flags        : {', '.join(job.get('red_flags', [])) or 'None'}",
        f"",
        f"KEYWORDS USED IN TAILORED RESUME",
        f"-" * 30,
    ]
    for kw in tailored.get("keywords_used", []):
        lines.append(f"  • {kw}")

    lines += [
        "",
        "CHANGES MADE TO RESUME",
        "-" * 30,
        tailored.get("changes_made", "N/A"),
        "",
        "SUGGESTED FOLLOW-UP MESSAGE",
        "-" * 30,
        f"Subject: Following Up — {job.get('title', '')} Application",
        "",
        f"Hi [Hiring Manager],",
        "",
        f"I wanted to follow up on my application for the {job.get('title', '')} "
        f"position at {job.get('company', '')}. I am very interested in this opportunity "
        f"and confident that my 20+ years of accounting and bookkeeping experience, "
        f"QBO/Xero certifications, and remote work background make me a strong fit.",
        "",
        f"I'd love the chance to connect and learn more about the role. "
        f"Please let me know if you need anything else from me.",
        "",
        f"Thank you,",
        f"Adriana Flores",
        f"+312-802-0811 | adriana@adrianaflores.com",
    ]

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    logger.info(f"Note saved: {output_path}")


# ── Main tailor function ──────────────────────────────────────────────────────

def safe_filename(text: str) -> str:
    """Sanitize a string for use in file names."""
    return re.sub(r"[^\w\-_]", "_", text).strip("_")[:40]


def tailor_for_job(
    job: Dict,
    profile: Dict,
    config: Dict,
    api_key: str,
    base_dir: str,
) -> Dict:
    """
    Full tailoring pipeline for one job:
      1. AI resume rewrite
      2. DOCX + PDF generation
      3. Application note
    Returns dict of output paths.
    """
    company = safe_filename(job.get("company", "Company"))
    title = safe_filename(job.get("title", "Role"))
    stem = f"Adriana_Flores_{company}_{title}"

    resumes_dir = os.path.join(base_dir, config["output"]["resumes_dir"])
    notes_dir = os.path.join(base_dir, config["output"]["notes_dir"])

    docx_path = os.path.join(resumes_dir, f"{stem}_Resume.docx")
    pdf_path = os.path.join(resumes_dir, f"{stem}_Resume.pdf")
    note_path = os.path.join(notes_dir, f"{stem}_Note.txt")

    model = config.get("api", {}).get("anthropic_model", "claude-sonnet-4-6")

    logger.info(f"Tailoring resume for: {job.get('title')} at {job.get('company')}")
    tailored = tailor_resume_with_ai(profile, job, model, api_key)

    generate_docx(profile, tailored, docx_path)
    generate_pdf(profile, tailored, pdf_path)
    generate_application_note(job, tailored, note_path)

    return {
        "docx": docx_path,
        "pdf": pdf_path,
        "note": note_path,
        "keywords_used": tailored.get("keywords_used", []),
        "changes_made": tailored.get("changes_made", ""),
    }
