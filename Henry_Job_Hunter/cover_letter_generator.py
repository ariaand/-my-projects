"""
cover_letter_generator.py — AI-powered cover letter creation.

Generates a warm, professional, concise cover letter tailored to each job.
Saves as DOCX and PDF.
"""

import json
import logging
import os
import re
from datetime import datetime
from typing import Dict

import anthropic
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors

logger = logging.getLogger(__name__)

COVER_LETTER_PROMPT = """You are an expert cover letter writer specializing in accounting and bookkeeping professionals.

Write a warm, professional, concise cover letter for the candidate below applying to this job.

RULES:
- 3 short paragraphs + closing (total ~250-300 words)
- Paragraph 1: Hook — why this specific role at this company interests her
- Paragraph 2: 2-3 specific strengths that match the job (use keywords from the JD; only real skills)
- Paragraph 3: Brief statement about remote work capability and bilingual value (only if relevant to JD)
- Closing: Brief, confident, call-to-action
- DO NOT fabricate experience, certifications, or skills not in the profile
- Tone: warm, confident, professional — not stiff or overly formal
- Do NOT use "I am writing to apply for..." or cliché openers

CANDIDATE PROFILE:
{profile}

JOB TITLE: {job_title}
COMPANY: {company}

JOB DESCRIPTION:
{job_description}

Return ONLY the letter body text (no subject line, no date, no address block).
Start directly with "Dear Hiring Team," or "Dear [Company] Team,"
End with:
Warm regards,
Adriana Flores
+312-802-0811 | adriana@adrianaflores.com
"""


def generate_cover_letter_text(
    profile: Dict,
    job: Dict,
    model: str,
    api_key: str,
) -> str:
    """Call Claude to generate the cover letter body."""
    client = anthropic.Anthropic(api_key=api_key)

    prompt = COVER_LETTER_PROMPT.format(
        profile=json.dumps(profile, indent=2),
        job_title=job.get("title", ""),
        company=job.get("company", ""),
        job_description=job.get("description", "")[:3000],
    )

    message = client.messages.create(
        model=model,
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}],
    )

    return message.content[0].text.strip()


def _save_cover_letter_docx(
    letter_text: str,
    profile: Dict,
    job: Dict,
    output_path: str,
) -> None:
    """Write cover letter to DOCX with clean business letter formatting."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    def para(text="", size=11, bold=False, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = bold
        return p

    # Header
    para(profile["name"].upper(), size=14, bold=True, space_after=2,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    contact = "  |  ".join(filter(None, [
        profile.get("phone", ""),
        profile.get("email", ""),
        profile.get("location", ""),
    ]))
    para(contact, size=9.5, space_after=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Date
    para(datetime.now().strftime("%B %d, %Y"), size=11, space_after=6)

    # Addressee block
    para(job.get("company", "Hiring Team"), size=11, space_after=2)
    para("Remote Position", size=11, space_after=12)

    # Letter body — split by newlines and render each paragraph
    for block in letter_text.split("\n\n"):
        block = block.strip()
        if block:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(block.replace("\n", " "))
            run.font.name = "Calibri"
            run.font.size = Pt(11)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    logger.info(f"Cover letter DOCX saved: {output_path}")


def _save_cover_letter_pdf(
    letter_text: str,
    profile: Dict,
    job: Dict,
    output_path: str,
) -> None:
    """Write cover letter to PDF."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=1.25 * inch,
        rightMargin=1.25 * inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    DARK = colors.HexColor("#222222")
    BLUE = colors.HexColor("#1F497D")

    name_style = ParagraphStyle("Name", fontSize=14, textColor=BLUE,
                                 alignment=1, spaceAfter=4, fontName="Helvetica-Bold")
    contact_style = ParagraphStyle("Contact", fontSize=9.5, textColor=DARK,
                                    alignment=1, spaceAfter=16)
    date_style = ParagraphStyle("Date", fontSize=11, textColor=DARK, spaceAfter=6)
    addr_style = ParagraphStyle("Addr", fontSize=11, textColor=DARK, spaceAfter=4)
    body_style = ParagraphStyle("Body", fontSize=11, textColor=DARK,
                                 leading=16, spaceAfter=10, alignment=TA_LEFT)

    story = []
    story.append(Paragraph(profile["name"].upper(), name_style))

    contact = "  |  ".join(filter(None, [
        profile.get("phone", ""),
        profile.get("email", ""),
        profile.get("location", ""),
    ]))
    story.append(Paragraph(contact, contact_style))
    story.append(Paragraph(datetime.now().strftime("%B %d, %Y"), date_style))
    story.append(Paragraph(job.get("company", "Hiring Team"), addr_style))
    story.append(Paragraph("Remote Position", addr_style))
    story.append(Spacer(1, 12))

    for block in letter_text.split("\n\n"):
        block = block.strip()
        if block:
            story.append(Paragraph(block.replace("\n", " "), body_style))

    doc.build(story)
    logger.info(f"Cover letter PDF saved: {output_path}")


def safe_filename(text: str) -> str:
    return re.sub(r"[^\w\-_]", "_", text).strip("_")[:40]


def generate_cover_letter(
    job: Dict,
    profile: Dict,
    config: Dict,
    api_key: str,
    base_dir: str,
) -> Dict:
    """
    Generate cover letter DOCX + PDF for a single job.
    Returns dict with file paths.
    """
    company = safe_filename(job.get("company", "Company"))
    title = safe_filename(job.get("title", "Role"))
    stem = f"Adriana_Flores_{company}_{title}"

    cl_dir = os.path.join(base_dir, config["output"]["cover_letters_dir"])
    docx_path = os.path.join(cl_dir, f"{stem}_CoverLetter.docx")
    pdf_path = os.path.join(cl_dir, f"{stem}_CoverLetter.pdf")

    model = config.get("api", {}).get("anthropic_model", "claude-sonnet-4-6")

    logger.info(f"Generating cover letter for: {job.get('title')} at {job.get('company')}")
    letter_text = generate_cover_letter_text(profile, job, model, api_key)

    _save_cover_letter_docx(letter_text, profile, job, docx_path)
    _save_cover_letter_pdf(letter_text, profile, job, pdf_path)

    return {"docx": docx_path, "pdf": pdf_path, "text": letter_text}
